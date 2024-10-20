from docx import Document
from docx.shared import Inches
from PIL import Image

def image_to_word(image_path, doc_path):
    doc = Document()
    try:
        img = Image.open(image_path)
        img.show()  
    except Exception as e:
        print(f"Error opening image: {e}")
        return
    doc.add_heading('Image to Word', 0)
    doc.add_picture(image_path, width=Inches(5.0))  
    doc.save(doc_path)

    print(f"Document saved successfully as {doc_path}")
image_path = "D:\Git1\zoro.jpg"
doc_path = "D:\Git1\created.docx"
image_to_word(image_path, doc_path)